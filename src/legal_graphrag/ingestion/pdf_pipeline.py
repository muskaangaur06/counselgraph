"""
PDF -> RAG ingestion pipeline for Google Colab.

Stages: extract text/tables with pdfplumber, OCR any low-text (scanned)
pages, clean and chunk the text with sliding-window overlap, turn tables
into their own markdown chunks, persist metadata to JSON, then embed
everything and store it in a local Chroma collection.

Run this in a Colab cell first (system packages + python packages):

    !apt-get -qq update && apt-get -qq install -y poppler-utils tesseract-ocr

    # --no-deps is important: Colab ships a specific torch/torchvision pair
    # already. Installing sentence-transformers normally lets pip silently
    # downgrade torch to satisfy its pin, which breaks torchvision (still
    # expecting the newer torch) and can crash embedding with
    # "RuntimeError: Numpy is not available". --no-deps keeps Colab's
    # existing torch/torchvision untouched; we then add sentence-transformers'
    # other dependencies explicitly, without letting them touch torch either.
    !pip -q install pdfplumber pdf2image pytesseract chromadb
    !pip -q install sentence-transformers --no-deps
    !pip -q install --upgrade-strategy only-if-needed \
        transformers tokenizers huggingface-hub safetensors scikit-learn scipy Pillow tqdm

    # Then: Runtime -> Restart session (required: a pip install that touches
    # any of torch's neighbors needs a fresh process, not just a fresh cell).

Then upload a PDF (or mount Drive) and run:

    from pdf_rag_pipeline import run_pipeline
    run_pipeline("data/uploads/my_document.pdf")   # or an absolute /content/... path in Colab
"""

from __future__ import annotations

import json
import os
import re
import uuid
from dataclasses import dataclass, asdict, field
from typing import Optional

import pdfplumber
from pdf2image import convert_from_path
import pytesseract
import docx as python_docx

from sentence_transformers import SentenceTransformer
import chromadb


# Config

MIN_CHARS_FOR_TEXT_PAGE = 20   # pages with fewer extracted characters than this are treated as "no text"
OCR_DPI = 300                  # resolution used when rasterizing a page for OCR: higher = better OCR, slower
CHUNK_SIZE_CHARS = 1000        # target characters per chunk
CHUNK_OVERLAP_CHARS = 150      # overlap between consecutive chunks, preserves context across chunk boundaries
EMBEDDING_MODEL_NAME = "sentence-transformers/all-mpnet-base-v2"

# Resolved relative to the project root (two levels up from src/legal_graphrag/ingestion/)
# so it works the same whether run from VS Code, a terminal, or Colab (with /content mounted).
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
CHROMA_PERSIST_DIR = os.path.join(_PROJECT_ROOT, "data", "chroma_db")
DEFAULT_METADATA_DIR = os.path.join(_PROJECT_ROOT, "data", "metadata")


# Data structures

@dataclass
class PageRecord:
    page_number: int              # 1-indexed
    text: str
    source: str                   # "pdfplumber" or "ocr"
    char_count: int
    section: Optional[str] = None  # heading detected on this page, if any


@dataclass
class TableRecord:
    table_id: str
    document_name: str
    page_number: int
    table_index: int              # index of this table within its page (0-based)
    section: Optional[str]
    markdown: str
    num_rows: int
    num_cols: int


@dataclass
class ChunkRecord:
    chunk_id: str
    document_name: str
    chunk_index: int
    text: str
    page_start: int
    page_end: int
    section: Optional[str]
    content_type: str = "text"                    # "text" or "table"
    sources: list = field(default_factory=list)    # e.g. ["pdfplumber"], ["ocr"], or both if it spans pages


# Step 1: Extract text AND tables with pdfplumber

def extract_page_content(pdf_path: str) -> tuple[list[PageRecord], dict[int, list[list[list[Optional[str]]]]]]:
    """Extract text and tables page-by-page with pdfplumber in one pass. Scanned
    pages come back with little text and no tables (step 2 handles the OCR side;
    scanned tables aren't recovered here, only their OCR'd text will be)."""
    pages: list[PageRecord] = []
    raw_tables: dict[int, list[list[list[Optional[str]]]]] = {}

    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            raw_text = page.extract_text() or ""
            pages.append(
                PageRecord(
                    page_number=i,
                    text=raw_text,
                    source="pdfplumber",
                    char_count=len(raw_text.strip()),
                )
            )

            # default line-detection works for ruled tables; borderless/whitespace
            # tables may need custom table_settings (see pdfplumber docs)
            tables_on_page = page.extract_tables()
            if tables_on_page:
                raw_tables[i] = tables_on_page

    return pages, raw_tables


def extract_docx_content(docx_path: str) -> tuple[list[PageRecord], dict[int, list[list[list[Optional[str]]]]]]:
    """DOCX has no fixed page boundaries, so the whole document is treated as page 1,
    downstream chunking still works the same way on the concatenated text."""
    document = python_docx.Document(docx_path)

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    pages = [PageRecord(page_number=1, text=full_text, source="docx", char_count=len(full_text.strip()))]

    raw_tables: dict[int, list[list[list[Optional[str]]]]] = {}
    tables_on_page = []
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows:
            tables_on_page.append(rows)
    if tables_on_page:
        raw_tables[1] = tables_on_page

    return pages, raw_tables


# Step 2: Detect pages with little or no text

def detect_low_text_pages(pages: list[PageRecord], threshold: int = MIN_CHARS_FOR_TEXT_PAGE) -> list[int]:
    """Return page numbers that fell below `threshold` chars: candidates for OCR."""
    return [p.page_number for p in pages if p.char_count < threshold]


# Step 3: OCR the low-text pages

def ocr_pages(pdf_path: str, page_numbers: list[int], dpi: int = OCR_DPI) -> dict[int, str]:
    """Rasterize only the given pages and run Tesseract OCR on each (OCR is slow, so skip pages that don't need it)."""
    if not page_numbers:
        return {}

    ocr_results: dict[int, str] = {}

    # rasterize one page at a time to keep memory bounded on large PDFs
    for page_number in page_numbers:
        images = convert_from_path(
            pdf_path, dpi=dpi, first_page=page_number, last_page=page_number
        )
        if not images:
            ocr_results[page_number] = ""
            continue

        ocr_text = pytesseract.image_to_string(images[0])
        ocr_results[page_number] = ocr_text

    return ocr_results


def merge_ocr_results(pages: list[PageRecord], ocr_results: dict[int, str]) -> list[PageRecord]:
    """Replace a page's text/source with its OCR output, where OCR was run."""
    merged = []
    for page in pages:
        if page.page_number in ocr_results:
            ocr_text = ocr_results[page.page_number]
            merged.append(
                PageRecord(
                    page_number=page.page_number,
                    text=ocr_text,
                    source="ocr",
                    char_count=len(ocr_text.strip()),
                )
            )
        else:
            merged.append(page)
    return merged


# Section detection (shared by text chunking and table records)

_HEADING_RE = re.compile(r"^[A-Z0-9][A-Za-z0-9 ,.:\-()]{2,80}$")


def detect_section_heading(text: str) -> Optional[str]:
    """Cheap heuristic: if the page's first non-empty line looks like a heading, use it as the section label."""
    for line in text.splitlines():
        candidate = line.strip()
        if not candidate:
            continue
        if len(candidate) <= 80 and _HEADING_RE.match(candidate):
            return candidate
        break  # only ever look at the first non-empty line
    return None


def compute_page_sections(pages: list[PageRecord]) -> dict[int, Optional[str]]:
    """Carry the most recent heading forward as each page's section, so text and tables on the same page agree."""
    page_sections: dict[int, Optional[str]] = {}
    cursor_section: Optional[str] = None
    for page in pages:
        heading = detect_section_heading(page.text)
        if heading:
            cursor_section = heading
        page_sections[page.page_number] = cursor_section
    return page_sections


# Step 4: Clean and chunk text

def clean_text(text: str) -> str:
    """Normalize whitespace and fix common PDF-extraction artifacts (dehyphenation, stray newlines, etc)."""
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)          # dehyphenate across line breaks
    text = re.sub(r"[ \t]+", " ", text)                     # collapse repeated spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n", text)                  # collapse 3+ blank lines to 1
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)             # single newlines -> space (keep paragraph breaks)
    text = "".join(ch for ch in text if ch.isprintable() or ch in "\n\t")
    return text.strip()


def _deterministic_chunk_id(document_name: str, chunk_index: int, text: str) -> str:
    """Derived from document_name + chunk_index + a content hash instead of a random
    UUID, so re-running ingestion for the same document upserts existing Chroma
    vectors rather than duplicating them (see db/dedup.py for the same idea applied
    to clauses)."""
    import hashlib
    import re as _re
    normalized = _re.sub(r"\s+", " ", text.strip().lower())
    content_hash = hashlib.sha256(normalized.encode("utf-8")).hexdigest()[:16]
    return f"{document_name}:{chunk_index}:{content_hash}"


def chunk_pages(pages: list[PageRecord], document_name: str,
                 page_sections: dict[int, Optional[str]],
                 chunk_size: int = CHUNK_SIZE_CHARS,
                 overlap: int = CHUNK_OVERLAP_CHARS,
                 start_chunk_index: int = 0) -> list[ChunkRecord]:
    """Cleans each page's text then chunks the concatenated document with a
    sliding window. Chunks can span pages but still track their page range/
    section for citations. Tables are chunked separately so rows never get sliced."""
    buffer_parts = []
    offset_to_page: list[tuple[int, int, str]] = []  # (start_offset, page_number, source)
    running_offset = 0

    for page in pages:
        cleaned = clean_text(page.text)
        offset_to_page.append((running_offset, page.page_number, page.source))
        buffer_parts.append(cleaned)
        running_offset += len(cleaned) + 1  # +1 for the joining space added below
        buffer_parts.append(" ")

    full_text = "".join(buffer_parts)

    def lookup(offset: int) -> tuple[int, str]:
        """Find the page/source active at a given character offset."""
        page_number, source = pages[0].page_number, pages[0].source
        for start, pnum, src in offset_to_page:
            if start <= offset:
                page_number, source = pnum, src
            else:
                break
        return page_number, source

    chunks: list[ChunkRecord] = []
    start = 0
    chunk_index = start_chunk_index
    text_len = len(full_text)

    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk_text = full_text[start:end].strip()

        if chunk_text:
            page_start, source_start = lookup(start)
            page_end, source_end = lookup(max(end - 1, start))
            sources = sorted({source_start, source_end})

            chunks.append(
                ChunkRecord(
                    chunk_id=_deterministic_chunk_id(document_name, chunk_index, chunk_text),
                    document_name=document_name,
                    chunk_index=chunk_index,
                    text=chunk_text,
                    page_start=page_start,
                    page_end=page_end,
                    section=page_sections.get(page_start),
                    content_type="text",
                    sources=sources,
                )
            )
            chunk_index += 1

        if end == text_len:
            break
        start = end - overlap  # step forward, re-including `overlap` characters of context

    return chunks


# Step 5: Turn extracted tables into their own chunks

def table_to_markdown(table: list[list[Optional[str]]]) -> str:
    """Render a pdfplumber raw table as a markdown table, so rows/columns survive instead of flattening into text."""
    if not table:
        return ""

    def clean_cell(cell: Optional[str]) -> str:
        return (cell or "").replace("\n", " ").replace("|", "/").strip()

    header, *body_rows = table
    header_line = "| " + " | ".join(clean_cell(c) for c in header) + " |"
    separator_line = "| " + " | ".join("---" for _ in header) + " |"
    body_lines = [
        "| " + " | ".join(clean_cell(c) for c in row) + " |"
        for row in body_rows
    ]
    return "\n".join([header_line, separator_line, *body_lines])


def build_table_records(raw_tables: dict[int, list[list[list[Optional[str]]]]],
                          document_name: str,
                          page_sections: dict[int, Optional[str]]) -> list[TableRecord]:
    """Convert every raw extracted table into a TableRecord with its markdown rendering."""
    records: list[TableRecord] = []
    for page_number, tables_on_page in sorted(raw_tables.items()):
        for table_index, raw_table in enumerate(tables_on_page):
            markdown = table_to_markdown(raw_table)
            if not markdown:
                continue
            records.append(
                TableRecord(
                    table_id=_deterministic_chunk_id(document_name, f"table-p{page_number}-{table_index}", markdown),
                    document_name=document_name,
                    page_number=page_number,
                    table_index=table_index,
                    section=page_sections.get(page_number),
                    markdown=markdown,
                    num_rows=len(raw_table),
                    num_cols=max((len(r) for r in raw_table), default=0),
                )
            )
    return records


def build_table_chunks(table_records: list[TableRecord], start_chunk_index: int = 0) -> list[ChunkRecord]:
    """Wrap each TableRecord as its own ChunkRecord so it flows through persistence/embedding like text chunks do."""
    chunks: list[ChunkRecord] = []
    for i, t in enumerate(table_records):
        chunks.append(
            ChunkRecord(
                chunk_id=t.table_id,
                document_name=t.document_name,
                chunk_index=start_chunk_index + i,
                text=t.markdown,
                page_start=t.page_number,
                page_end=t.page_number,
                section=t.section,
                content_type="table",
                sources=["pdfplumber"],
            )
        )
    return chunks


# Step 6: Persist metadata for traceability

def persist_metadata(chunks: list[ChunkRecord], output_path: str) -> str:
    """Write every chunk's metadata (and text) to a JSON file for later auditing, independent of the vector store."""
    payload = [asdict(c) for c in chunks]
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return output_path


# Step 7: Embed and store in a vector database

def embed_and_store(chunks: list[ChunkRecord], collection_name: str,
                     embedding_model_name: str = EMBEDDING_MODEL_NAME,
                     persist_dir: str = CHROMA_PERSIST_DIR,
                     extra_metadata: Optional[dict] = None):
    """Embeds every chunk and upserts it into a persistent Chroma collection.
    extra_metadata (contract_type, parties, etc) gets merged into every chunk's
    metadata since Chroma filters per-chunk, not per-document."""
    extra_metadata = extra_metadata or {}
    embedder = SentenceTransformer(embedding_model_name)
    texts = [c.text for c in chunks]
    # convert_to_tensor + .tolist() sidesteps torch's numpy bridge, which
    # sometimes breaks in Colab with "RuntimeError: Numpy is not available"
    embeddings = embedder.encode(
        texts, normalize_embeddings=True, show_progress_bar=True, convert_to_tensor=True
    ).tolist()

    client = chromadb.PersistentClient(path=persist_dir)
    collection = client.get_or_create_collection(name=collection_name)

    collection.upsert(
        ids=[c.chunk_id for c in chunks],
        embeddings=embeddings,
        documents=texts,
        metadatas=[
            {
                "document_name": c.document_name,
                "chunk_index": c.chunk_index,
                "page_start": c.page_start,
                "page_end": c.page_end,
                "section": c.section or "",
                "content_type": c.content_type,
                "sources": ",".join(c.sources),
                **extra_metadata,
            }
            for c in chunks
        ],
    )

    return collection, embedder


def query_collection(collection, embedder: SentenceTransformer, query: str, top_k: int = 5):
    """Embed a query and retrieve the top_k most similar chunks, with citations."""
    query_embedding = embedder.encode([query], normalize_embeddings=True).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=top_k)

    hits = []
    for doc, meta, dist in zip(results["documents"][0], results["metadatas"][0], results["distances"][0]):
        hits.append({"text": doc, "metadata": meta, "distance": dist})
    return hits


# Orchestration

def run_pipeline(pdf_path: str, collection_name: Optional[str] = None,
                  metadata_output_path: Optional[str] = None,
                  sample_query: Optional[str] = None) -> dict:
    """Run the full extract -> OCR -> chunk -> table -> persist -> embed pipeline for one PDF."""
    document_name = os.path.basename(pdf_path)
    collection_name = collection_name or re.sub(r"\W+", "_", document_name)
    metadata_output_path = metadata_output_path or os.path.join(DEFAULT_METADATA_DIR, f"{document_name}.metadata.json")

    print(f"[1/7] Extracting text and tables with pdfplumber: {document_name}")
    pages, raw_tables = extract_page_content(pdf_path)
    num_raw_tables = sum(len(t) for t in raw_tables.values())
    print(f"       {num_raw_tables} table(s) found across {len(raw_tables)} page(s)")

    print("[2/7] Detecting low-text pages")
    low_text_pages = detect_low_text_pages(pages)
    print(f"       {len(low_text_pages)} of {len(pages)} pages need OCR: {low_text_pages}")

    print("[3/7] Running OCR on low-text pages")
    ocr_results = ocr_pages(pdf_path, low_text_pages)
    pages = merge_ocr_results(pages, ocr_results)

    print("[4/7] Cleaning and chunking text")
    page_sections = compute_page_sections(pages)
    text_chunks = chunk_pages(pages, document_name, page_sections)
    print(f"       {len(text_chunks)} text chunks produced")

    print("[5/7] Building table chunks")
    table_records = build_table_records(raw_tables, document_name, page_sections)
    table_chunks = build_table_chunks(table_records, start_chunk_index=len(text_chunks))
    print(f"       {len(table_chunks)} table chunks produced")

    chunks = text_chunks + table_chunks

    print("[6/7] Persisting page/section/source/table metadata")
    metadata_path = persist_metadata(chunks, metadata_output_path)
    print(f"       written to {metadata_path}")

    print("[7/7] Embedding and storing in Chroma")
    collection, embedder = embed_and_store(chunks, collection_name)
    print(f"       stored in collection '{collection_name}'")

    result = {
        "document_name": document_name,
        "collection_name": collection_name,
        "metadata_path": metadata_path,
        "num_pages": len(pages),
        "num_ocr_pages": len(low_text_pages),
        "num_text_chunks": len(text_chunks),
        "num_tables": len(table_records),
        "num_chunks": len(chunks),
    }

    if sample_query:
        hits = query_collection(collection, embedder, sample_query)
        result["sample_query"] = sample_query
        result["sample_results"] = hits
        print(f"\nSample query: {sample_query!r}")
        for h in hits:
            m = h["metadata"]
            tag = "[TABLE]" if m["content_type"] == "table" else "[TEXT] "
            print(f"  {tag} p.{m['page_start']}-{m['page_end']} [{m['section'] or 'no section'}] "
                  f"(dist={h['distance']:.4f}): {h['text'][:120]}...")

    return result


if __name__ == "__main__":
    # Example (adjust the path to a PDF uploaded in your Colab session):
    run_pipeline("data/uploads/sample.pdf", sample_query="What is the termination clause?")
